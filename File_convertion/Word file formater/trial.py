import json
from collections import Counter
from statistics import mean
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_general_formatting(input_docx, json_file):
    doc = Document(input_docx)

    # Data structures to collect formatting attributes
    font_sizes = []
    alignments = []
    line_spacings = []
    space_before_values = []
    space_after_values = []
    left_indents = []
    right_indents = []
    first_line_indents = []

    for para in doc.paragraphs:
        if para.alignment is not None:
            alignments.append(para.alignment)
        if para.paragraph_format.line_spacing is not None:
            line_spacings.append(para.paragraph_format.line_spacing)
        if para.paragraph_format.space_before is not None:
            space_before_values.append(para.paragraph_format.space_before.pt)
        if para.paragraph_format.space_after is not None:
            space_after_values.append(para.paragraph_format.space_after.pt)
        if para.paragraph_format.left_indent is not None:
            left_indents.append(para.paragraph_format.left_indent.pt)
        if para.paragraph_format.right_indent is not None:
            right_indents.append(para.paragraph_format.right_indent.pt)
        if para.paragraph_format.first_line_indent is not None:
            first_line_indents.append(para.paragraph_format.first_line_indent.pt)
        for run in para.runs:
            if run.font.size is not None:
                font_sizes.append(run.font.size.pt)

    # Compute general formatting profile
    general_formatting = {
        "font_size": round(mean(font_sizes), 2) if font_sizes else None,
        "alignment": Counter(alignments).most_common(1)[0][0] if alignments else None, 
        "line_spacing": mean(line_spacings) if line_spacings else None,
        "space_before": mean(space_before_values) if space_before_values else None,
        "space_after": mean(space_after_values) if space_after_values else None,
        "left_indent": mean(left_indents) if left_indents else None,
        "right_indent": mean(right_indents) if right_indents else None,
        "first_line_indent": mean(first_line_indents) if first_line_indents else None
    }

    # Save general formatting to a JSON file
    with open(json_file, 'w') as f:
        json.dump(general_formatting, f, indent=4)


def apply_general_formatting(json_file, input_docx, output_docx):
    # Load general formatting from JSON file
    with open(json_file, 'r') as f:
        general_formatting = json.load(f)

    # Load the input document
    doc = Document(input_docx)

    for para in doc.paragraphs:
        # Apply paragraph-level formatting
        if general_formatting["alignment"] is not None:
            para.alignment = WD_ALIGN_PARAGRAPH(general_formatting["alignment"])

        if general_formatting["line_spacing"]:
            para.paragraph_format.line_spacing = general_formatting["line_spacing"]
        if general_formatting["space_before"]:
            para.paragraph_format.space_before = Pt(general_formatting["space_before"])
        if general_formatting["space_after"]:
            para.paragraph_format.space_after = Pt(general_formatting["space_after"])
        if general_formatting["left_indent"]:
            para.paragraph_format.left_indent = Pt(general_formatting["left_indent"])
        if general_formatting["right_indent"]:
            para.paragraph_format.right_indent = Pt(general_formatting["right_indent"])
        if general_formatting["first_line_indent"]:
            para.paragraph_format.first_line_indent = Pt(general_formatting["first_line_indent"])

        for run in para.runs:
            # Apply run-level (text-level) formatting
            if general_formatting["font_size"]:
                run.font.size = Pt(general_formatting["font_size"])

    # Save the modified document
    doc.save(output_docx)


if __name__ == "__main__":
    # Extract general formatting from the first Word document
    extract_general_formatting(r'E:\sync\Research\College details\canvas\Idea-or-Proof-of-Concept-Submisson-Format.docx', 'general_formatting.json')

    # Apply general formatting to the second Word document
    # apply_general_formatting('general_formatting.json', 'input2.docx', 'output.docx')
